import docx 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
  
originalList = Document('<path to file>/05-11-24 Prayer List.docx')
newList = docx.Document()

newList_para = newList.add_paragraph('Prayer List')
newList_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
newList_para.paragraph_format.space_after = Inches(0)

search = "Pray for our Pastor"
print('\nThe whole content of the document:->>>\n')
p = 0
for para in originalList.paragraphs: 
    print(para.text)
    print("line Paragraph #", p ,":")
    
    if(p > 2 and p <= 4):
        newList_para = newList.add_paragraph(originalList.paragraphs[p].text)
        newList_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        newList_para.paragraph_format.space_after = Inches(0)
    
    if(p > 4):
        newList_para = newList.add_paragraph(originalList.paragraphs[p].text)
        newList_para.paragraph_format.space_after = Inches(0)
    line = originalList.paragraphs[p+1].text
    if((p > 2) and (search in line)):
        break
    p+=1
newList.save("<path to file>/Prayer List.docx")

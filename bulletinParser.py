import docx 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import doc2docx

directoryLocation = "C:/Users/jmswe/OneDrive/Documents/1Coding/Python/Church scripts/"
prayerlist = "05-11-24 Prayer List.doc"
finalPrayerList = "05-11-24 Prayer List final.docx"
def parse():
    originalList = Document(directoryLocation+finalPrayerList)
    newList = docx.Document()
    newList_para = newList.add_paragraph('Prayer List')
    newList_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    newList_para.paragraph_format.space_after = Inches(0)
    search = "Pray for our Pastor"
    p = 0
    for para in originalList.paragraphs: 
        if(p > 2 and p <= 4):
            newList_para = newList.add_paragraph(originalList.paragraphs[p].text)
            newList_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            newList_para.paragraph_format.space_after = Inches(0)
        if(p > 4):
            newList_para = newList.add_paragraph(originalList.paragraphs[p].text)
            newList_para.paragraph_format.space_after = Inches(0)
        line = originalList.paragraphs[p+1].text
        if((p > 2) and (search in line)):
            break
        p+=1
    newList.save("C:/Users/jmswe/OneDrive/Documents/1Coding/Python/Church scripts/Prayer List.docx")

print('doc2docx Status: starting')
doc2docx.parse( directoryLocation+prayerlist, directoryLocation+finalPrayerList)
print('doc2docx Status: Done')
print('parsing Status: starting')
parse()
print('parsing Status: Done')

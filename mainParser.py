import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import doc2docx
import parseBulletin

directoryLocation = "<directory to prayer lit>/"
prayerlist = "05-11-20 Prayer List.doc"
docxPrayerList = "05-11-20 Prayer List final.docx"
finalPrayerList = "Prayer List.docx"
newBulletinLoc = "<path to file>/Bulletin.docx"
originalBulletinLoc = "<path to file>/05-11-20 - Bulletin - Transfers.docx"
email1 = "johndoe@hotmail.com "
email2 = " Email janedoe@hotmail.com.. Text ###-###-####."
def parsePayerList():
    originalList = Document(directoryLocation+docxPrayerList)
    newList = docx.Document()
    newList_para = newList.add_paragraph('Prayer List')
    newList_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    newList_para.paragraph_format.space_after = Inches(0)
    search = "Pray for our Pastor"
    p = 0
    for para in originalList.paragraphs:
        if(p > 2 and p <= 4):
            newList_para = newList.add_paragraph(originalList.paragraphs[p].text)
            newList_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            newList_para.paragraph_format.space_after = Inches(0)
        if(p > 4):
            newList_para = newList.add_paragraph(originalList.paragraphs[p].text)
            newList_para.paragraph_format.space_after = Inches(0)
        if((p > 2) and (search in originalList.paragraphs[p].text)):
            break
        p+=1
    newList.save(directoryLocation+finalPrayerList)

print('doc2docx PrayerList Status: starting')
doc2docx.parse( directoryLocation+prayerlist, directoryLocation+docxPrayerList)
print('doc2docx Status: Done')
print('parsing PrayerList Status: starting')
parsePayerList()
print('parsing PrayerListStatus: Done')
print('parsing Bulletin Status: starting')
parseBulletin.bulletinParsing(email1, email2, originalBulletinLoc, newBulletinLoc)
print('parsing Bulletin Status: Done')
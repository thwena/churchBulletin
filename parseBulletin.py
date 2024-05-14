import docx 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches
# create an instance of a  
# word document we want to open 



def stripsBlanks(string):
    string = string.strip()
    array = [char for char in string]
    flag = False
    if(string == ""):
        return string
    for i in range(len(array)):
        
        if(array[i] == " ") or (array[i] == "\t"):
            if( ((array[i-1] == " ") or (array[i-1] == "\t"))):
                array[i] = ""
                flag = True
        if(array[i] == "$"):
            array[i-1] = "\t"
            flag = True
    if flag :
        temp = ""
        for i in array:
            temp += i
        return temp
    return string

def upcoming(string):
    array = [char for char in string]
    flag = False
    if(string == ""):
        return string
    for i in range(len(array)):
        if(array[i] == "\t"):
            if( (array[i-1] == "\t") ):
                array[i-1] = ""
                flag = True
    if flag :
        temp = ""
        for i in array:
            temp += i
        return temp
    return string
        
def bulletinParsing(email1, email2, originalBulletinLoc, newBulletinLoc):
    newBulletin = docx.Document()
    originalBulletin = Document(originalBulletinLoc) 
    flag1 = -1
    flag2 = -1
    flag3 = -1
    flag4 = -1
    p=0
    
    for para in originalBulletin.paragraphs:
        if("UPCOMING" in originalBulletin.paragraphs[p].text):
            flag1 = p
        if("PATH" in originalBulletin.paragraphs[p].text):
            flag2 = 0
        if("Benediction" in originalBulletin.paragraphs[p].text):
            flag3 = 0
        if( p < 6):
            newBulletin_para = newBulletin.add_paragraph(stripsBlanks(originalBulletin.paragraphs[p].text))
            newBulletin_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            newBulletin_para.paragraph_format.space_after = Inches(0)
        elif( (p == 6)):
            newBulletin_para = newBulletin.add_paragraph(stripsBlanks(originalBulletin.paragraphs[p].text))
            tab_stop = newBulletin_para.paragraph_format.tab_stops
            tab_stop.add_tab_stop(Inches(5), WD_TAB_ALIGNMENT.RIGHT)
            tab_stop.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
            if(p == 6):
                newBulletin_para.paragraph_format.space_after = Inches(0)
        elif( p == 7 ):
            paragraph = '\t',stripsBlanks(originalBulletin.paragraphs[p].text)
            newBulletin_para = newBulletin.add_paragraph(paragraph)
            tab_stop = newBulletin_para.paragraph_format.tab_stops
            tab_stop.add_tab_stop(Inches(1.3), WD_TAB_ALIGNMENT.LEFT)
            tab_stop.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
            newBulletin_para.paragraph_format.space_after = Inches(0)
            newBulletin_para = newBulletin.add_paragraph("")
        elif((p >= 9) and ((flag3 == -1) or ("Benediction" in originalBulletin.paragraphs[p].text)) ):
            testString = originalBulletin.paragraphs[p].text.strip()
            if((testString != "")):
                if("Sermon" in originalBulletin.paragraphs[p].text):
                    newBulletin_para = newBulletin.add_paragraph((originalBulletin.paragraphs[p].text))
                    tab_stop = newBulletin_para.paragraph_format.tab_stops
                    tab_stop.add_tab_stop(Inches(1), WD_TAB_ALIGNMENT.CENTER)
                    tab_stop.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.CENTER)
                    tab_stop.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                    newBulletin_para.paragraph_format.space_after = Inches(.1)
                else:
                    newBulletin_para = newBulletin.add_paragraph((originalBulletin.paragraphs[p].text))
                    tab_stop = newBulletin_para.paragraph_format.tab_stops
                    tab_stop.add_tab_stop(Inches(3), WD_TAB_ALIGNMENT.CENTER)
                    tab_stop.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                    newBulletin_para.paragraph_format.space_after = Inches(.1)
        elif( (p > flag1) and (flag2 != 0)):
            newBulletin_para = newBulletin.add_paragraph(upcoming(originalBulletin.paragraphs[p].text))
            tab_stop = newBulletin_para.paragraph_format.tab_stops
            tab_stop.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
            newBulletin_para.paragraph_format.space_after = Inches(0)
        elif(flag3 == 0):
            if(email1 in originalBulletin.paragraphs[p].text):
                newBulletin_para = newBulletin.add_paragraph((originalBulletin.paragraphs[p].text.replace(email1, "")))
                newBulletin_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                newBulletin_para.paragraph_format.space_after = Inches(.05)
            elif(email2 in originalBulletin.paragraphs[p].text):
                newBulletin_para = newBulletin.add_paragraph((originalBulletin.paragraphs[p].text.replace(email2, "")))
                newBulletin_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                newBulletin_para.paragraph_format.space_after = Inches(.05)
            else:
                newBulletin_para = newBulletin.add_paragraph(originalBulletin.paragraphs[p].text)
                newBulletin_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                newBulletin_para.paragraph_format.space_after = Inches(.05)
        p+=1
    newBulletin.save(newBulletinLoc) 
